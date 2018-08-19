# -*- coding: utf-8 -*-


import datetime
import io
import os
import zipfile
from collections import defaultdict
from lxml import etree as et
from repo import get_repository_dict
from trac.util.datefmt import format_date
from trac.util.text import to_unicode, to_utf8
from trac.versioncontrol.api import RepositoryManager
from trac.versioncontrol.web_ui.util import get_existing_node
from .model import PeerReviewModel, PeerReviewerModel, ReviewCommentModel, ReviewFileModel
from peerReviewPerform import file_data_from_repo

try:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor
    docx_support = True
except ImportError:
    docx_support = False

__author__ = 'Cinc'
__copyright__ = "Copyright 2016"
__license__ = "BSD"


def ensure_paragraph_styles(doc):
    """Add paragraph styles to the given document if they are not already defined.

    The user may use a template document for a report. Different parts of the added content
    use different styles. If the template doesn't define them defaults for the styles are used
    which are added here.

    The following styles are used:
    * 'Code': for listings of source files
    * 'Reviewcomment': review comment text
    * 'Reviewcommentinfo': info text for a comment. Author, date, id, ...
    """
    def have_paragraph_style(style_name):
        styles = doc.styles
        paragraph_styles = [s for s in styles if s.type == WD_STYLE_TYPE.PARAGRAPH]
        for style in paragraph_styles:
            if style.name == style_name:
                return True
        return False
    def add_style(style_data):
        style = doc.styles.add_style(style_data['name'], WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles['Normal']

        font_data = style_data['font']
        if font_data:
            font = style.font
            font.name = font_data[0]
            font.size = Pt(font_data[1])
            if font_data[2]:
                font.color.rgb = RGBColor(font_data[2][0], font_data[2][1], font_data[2][2])

    styles = [{'name': u'Code', 'font': [u'Courier New', 8, []]},
              {'name': u'Reviewcomment', 'font': [u'Arial', 10, []]},
              {'name': u'Reviewcommentinfo', 'font': [u'Arial', 8, [0x80, 0x80, 0x80]]}]

    for style_data in styles:
        if not have_paragraph_style(style_data['name']):
            add_style(style_data)


def add_review_info_to_table(env, doc, review_id):
    """Add info about the review to the given document.

    @param env: Trac envireonment object
    @param doc: python-docx document
    @param review_id: id of the review

    This function searches the document for a two column table with a marker field to add the information.
    If the table can't be found a heading 'Review Info' is added to the end of the document followed by
    a default table.

    Table format:

    |  Foo   |     Bar      |   Some header row
    |  Name  | $REVIEWNAME$ |   <- Marker field
    | Status |   $STATUS$   |
    |  ....  |    $....$    |   More rows, see cell_data below

    """
    def get_review_info(env, review_id):
        """Get a PeerReviewModel for the given review id and prepare some additional data used by the template"""
        review = PeerReviewModel(env, review_id)
        # review.html_notes = format_to_html(env, Context.from_request(req), review['notes'])
        review.date = format_date(review['created'], 'iso8601')
        return review

    def get_review_info_table(doc):
        for table in doc.tables:
            for row in table.rows:
                if len(row.cells) > 1 and row.cells[1].text == u'$REVIEWNAME$':
                    return table

        # Table not found, add it. This may be an empty template
        cell_data = [['Name', '$REVIEWNAME$'],
                     ['Status', '$STATUS$'],
                     ['ID', '$ID$'],
                     ['Project', '$PROJECT$'],
                     ['Author', '$AUTHOR$'],
                     ['Date', '$DATE$'],
                     ['Followup from', '$FOLLOWUP$']
                     ]
        doc.add_heading('Review Info', level=1)
        tbl = doc.add_table(len(cell_data), 2)
        for idx, data in enumerate(cell_data):
            tbl.rows[idx].cells[0].text = data[0]
            tbl.rows[idx].cells[1].text = data[1]
        return tbl

    rev_info = get_review_info(env, review_id)
    if rev_info:
        table = get_review_info_table(doc)
        if table:
            for row in table.rows:
                cell = row.cells[1]
                if cell.text == u'$REVIEWNAME$':
                    cell.text = rev_info['name']
                elif cell.text == u'$STATUS$':
                    cell.text = rev_info['status']
                elif cell.text == u'$PROJECT$':
                    cell.text = rev_info['project'] or ''
                elif cell.text == u'$AUTHOR$':
                    cell.text = rev_info['owner']
                elif cell.text == u'$DATE$':
                    cell.text = format_date(rev_info['created'], 'iso8601')
                elif cell.text == u'$ID$':
                    cell.text = str(rev_info['review_id'])
                elif cell.text == u'$FOLLOWUP$':
                    cell.text = str(rev_info['parent_id']) if rev_info['parent_id'] > 0  else '---'

def add_reviewers_to_table(env, doc, review_id):
    """Add reviewer names to a table in the document.

    The table will be searched by looking for a table row with $REVIEWER$ in the first column.
    The first reviewer will be inserted in the found column. For each subsequent reviewer a row is appended
    to the table. This means the column in question must be in the last row of the table.

    The table must have at least two columns:
    | Reviewer | Status | Some more columns... |

    """
    def get_reviewer_table(doc):
        """Find the table for inserting the reviewer names.

        This table must have two columns and at least one row with a cell in column 1 holding
        the text $REVIEWER$.
        First reviewer name will be inserted there and the following appended starting with
        the row/column holding the text. Thus the row in question must be the last one in the table.
        """
        for table in doc.tables:
            row = table.rows[-1]
            if row.cells[0].text == u'$REVIEWER$':
                return table, row

        # Table not found, add it. This may be an empty template
        cell_data = [['Reviewer', 'Status'],
                     ['$REVIEWER$', '']
                     ]
        doc.add_heading('Reviewer Info', level=1)
        tbl = doc.add_table(len(cell_data), 2)
        for idx, data in enumerate(cell_data):
            tbl.rows[idx].cells[0].text = data[0]
            tbl.rows[idx].cells[1].text = data[1]
        return tbl, tbl.rows[1]

    def get_reviewers(env, review_id):
        rm = PeerReviewerModel(env)
        rm.clear_props()
        rm['review_id'] = review_id
        return list(rm.list_matching_objects())

    reviewers = get_reviewers(env, review_id)
    table, row = get_reviewer_table(doc)
    if table:
        try:
            row.cells[0].text = reviewers[0]['reviewer']
            row.cells[1].text = reviewers[0]['status']
        except IndexError:
            pass
        for idx, reviewer in enumerate(reviewers[1:]):
            cells = table.add_row().cells
            try:
                cells[0].text = reviewer['reviewer']
                cells[1].text = reviewer['status']
            except IndexError:
                pass

def get_file_info(env, review_id):
    """Get a list of objects holding file and comment information for a review.

    Comment data is added to each of the files associated with the review.
    * Each file has an attribute 'comments'.
    * file_obj['comments'] holds a comment dict. Key: line number, value: ReviewCommentModel object

    @param env: Trac environment object
    @param review_id: review id
    @return: list of ReviewFileModel objects. This one has an additional attribute 'comments' which holds
             a dict with key: line number, value: list of comments for that line
    """
    def get_files_for_review_id(review_id):
        """Get all files belonging to the given review id. Provide the number of comments if asked for."""
        rfm = ReviewFileModel(env)
        rfm.clear_props()
        rfm['review_id'] = review_id
        return list(rfm.list_matching_objects())

    def get_comment_data_for_file(file_id):
        rcm = ReviewCommentModel(env)
        rcm.clear_props()
        rcm['file_id'] = file_id

        comments = list(rcm.list_matching_objects())

        the_dict = defaultdict(list)  # key: line_num, value: list of comments
        for comm in comments:
            comm['children'] = {}
            the_dict[comm['line_num']].append(comm)
        return the_dict

    items = get_files_for_review_id(review_id)
    for review_file in items:
        comments = get_comment_data_for_file(review_file['file_id'])
        review_file['comments'] = comments
    return items


def add_file_info_to_table(env, doc, review_id, file_info):
    """Find or create file information table and add file info.

    @param doc: python-docx Document object
    @param review_id: id of this review
    @param file_info: list of PeerReviewFile objects
    @return: None
    :param env:
    """
    def get_file_table(doc):
        for table in doc.tables:
            row = table.rows[-1]
            if len(row.cells) > 1 and row.cells[1].text == u'$FILEPATH$':
                return table

        # Table not found, add it. This may be an empty template
        cell_data = [['ID', 'Path', 'Hash', 'Revision', 'Comments', 'Status'],
                     ['', '$FILEPATH$', '', '', '', '']
                     ]
        doc.add_heading(u'Files', level=1)
        tbl = doc.add_table(len(cell_data), len(cell_data[0]))
        for idx, data in enumerate(cell_data):
            for i, val in enumerate(data):
                tbl.rows[idx].cells[i].text = val
        return tbl

    def get_num_comments(comments):
        """Count total number of comments for a file given a dict with comment information.

        @param comments: dict with key: line number, value: comments for this line
        @return: total number of comments
        """
        num = 0
        for value in comments.items():
            num += len(value)
        return num

    repodict = get_repository_dict(env)

    table = get_file_table(doc)
    if file_info and table:
        cells = table.rows[-1].cells

        for idx, item in enumerate(file_info):
            try:
                try:
                    prefix = repodict[item['repo']]['url'].rstrip('/')
                except KeyError:
                    prefix = ''
                if idx > 0:
                    cells = table.add_row().cells
                cells[0].text = str(item['file_id'])
                cells[1].text = prefix + item['path']
                cells[2].text = item['hash'] or ''
                cells[3].text = item['revision']
                cells[4].text = str(get_num_comments(item['comments']))
                cells[5].text = item['status'] or ''
            except IndexError:  # May happen if the template misses some table columns
                pass

def get_file_data(env, f_info):
    """Get file content from repository.

    @param env: Trac environment object
    @param f_info: PeerReviewFile object
    @return: file data for file specified by f_info
    """
    f_data = []
    repos = RepositoryManager(env).get_repository(f_info['repo'] or '')
    if not repos:
        return f_data

    rev = f_info['revision']
    if rev:
        rev = repos.normalize_rev(rev)
    rev_or_latest = rev or repos.youngest_rev
    node = get_existing_node(env, repos, f_info['path'], rev_or_latest)
    return file_data_from_repo(node)


def create_comment_tree(comments):
    comms = {}
    for c in comments:
        comms[c['comment_id']] = c

    all_keys = comms.keys()
    for key in all_keys:
        c = comms[key]
        if c['parent_id'] != -1 and c['parent_id'] in comms and c['parent_id'] != c['comment_id']:
            children_dict = comms[c['parent_id']]['children']
            children_dict[c['comment_id']] = c

    # Remove all comments without parent from root. These are still referenced in children dicts of some parent.
    for key in all_keys:
        c = comms[key]
        if c['parent_id'] != -1:
            del comms[key]

    return comms


def print_comment(par, comment, indent=0):
    """Add the given comment to the docx file.

    @param par: pythond-docx Paragraph. Comment will be inserted right before it.
    @param comment: ReviewComment object
    @param indent: number of tabs used for indenting
    @return: None
    """
    header = u"ID: %s: \t%s,\tAutor: %s" % (comment['comment_id'], format_date(comment['created']),
                                            comment['author'])
    par.insert_paragraph_before(u"\t"*indent + header, style=u"Reviewcommentinfo")
    par.insert_paragraph_before(u"\t"*indent + comment['comment'], style=u"Reviewcomment")
    children = comment['children']
    items = [c for id, c in children.items()]
    items = sorted(items, key=lambda item: item['comment_id'])
    for c in items:
        print_comment(par, c, indent + 1)


def add_file_data(env, doc, file_info):
    """Add file content and associated comments to document.

    The content of the given file is added to the document using the paragraph style 'Code'.
    The position in the document is specified by the string $FILECONTENT$ which must be in the
    template. If this string can't be found the data is appended to the end of the document.
    Note that a header with style 'Heading 2' will be added with the file path.

    @param env: Trac environment object
    @param doc: python-docx Document object
    @param file_info: PeerReviewFile object holding information about a file
    @return: None
    """
    def get_fileinfo_paragraph(doc):
        for par in doc.paragraphs:
            if u"$FILECONTENT$" in par.text:
                return par
        par = doc.add_paragraph(u"$FILECONTENT$")
        return par

    par = get_fileinfo_paragraph(doc)
    if par:
        for item in file_info:
            comments = item['comments']
            par.insert_paragraph_before(item['path'], style=u"Heading 2")
            f_data = get_file_data(env, item)
            line_nr = 0
            for line in f_data:
                line_nr += 1
                par.insert_paragraph_before("%s: %s" % (str(line_nr).rjust(4, ' '), to_unicode(line)),
                                            style=u'Code')
                if line_nr in comments:
                    par.insert_paragraph_before()
                    comm_tree = create_comment_tree(comments[line_nr])
                    items = [c for id, c in comm_tree.items()]
                    items = sorted(items, key=lambda item: item['comment_id'])
                    for comment in items:
                        print_comment(par, comment)
                    par.insert_paragraph_before()

        par.text = u""  # Overwrite marker text


def set_custom_doc_properties(zin, review):
    def set_element_txt(elm, txt):
        e = dom.xpath("//p:property[@name='%s']" % elm,
                      namespaces={'p': 'http://schemas.openxmlformats.org/officeDocument/2006/custom-properties'})
        if e:
            try:
                e[0][0].text = txt
            except IndexError:
                pass

    dom = et.fromstring(zin.read("docProps/custom.xml"))

    set_element_txt(u'VersionDate', datetime.datetime.today().strftime("%d.%m.%Y"))
    # Set document id
    set_element_txt(u'Dokumentnummer', '0')
    set_element_txt(u'MCNummer', review['project'])
    set_element_txt(u'VersionMajor', u'1')
    set_element_txt(u'VersionMinor', u'0')

    return et.tostring(dom)


def set_core_properties(doc, data):
    from datetime import datetime

    props = doc.core_properties
    props.created = datetime.now()


def set_core_doc_properties(zin, data):

    review = data['review']

    dom = et.fromstring(zin.read("docProps/core.xml"))

    e = dom.find("{http://purl.org/dc/elements/1.1/}creator")
    if e is not None:
        e.text = review['owner']
    e = dom.find("{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}lastModifiedBy")
    if e is not None:
        e.text = review['owner']

    e = dom.find("{http://purl.org/dc/elements/1.1/}subject")
    if e is not None:
        e.text = data['subject']

    e = dom.find("{http://purl.org/dc/elements/1.1/}title")
    if e is not None:
        e.text = data['title']

    e = dom.find("{http://schemas.openxmlformats.org/package/2006/metadata/core-properties}revision")
    if e is not None:
        e.text = '1'

    return et.tostring(dom)


def create_docx_for_review(env, data, template):
    """

    :param env: Trac environment object
    :param data: dictionary with information about the review
    :param template: path to docx template
    :return: docx file data
    """
    def template_exists(tpath):
        if not tpath:
            return False
        if os.path.isfile(tpath):
            return True
        return False

    review_id = data['review_id']
    if template and template_exists(template):
        doc = Document(template)
    else:
        doc = Document()

    ensure_paragraph_styles(doc)

    add_review_info_to_table(env, doc, review_id)
    add_reviewers_to_table(env, doc, review_id)

    file_info = get_file_info(env, review_id)
    add_file_info_to_table(env, doc, review_id, file_info)
    add_file_data(env, doc, file_info)

    set_core_properties(doc, data)
    buff = io.BytesIO()
    doc.save(buff)

    # Change custom properties
    out_buff = io.BytesIO()

    zin = zipfile.ZipFile(buff, 'r')
    with zipfile.ZipFile(out_buff, 'w') as zout:
        for item in zin.infolist():
            buf = zin.read(item.filename)
            if item.filename == 'docProps/custom.xml':
                zout.writestr("docProps/custom.xml",
                              set_custom_doc_properties(zin, data['review']))
            elif item.filename == 'docProps/core.xml':
                zout.writestr("docProps/core.xml",
                              set_core_doc_properties(zin, data))
            else:
                zout.writestr(item, buf)

    return out_buff.getvalue()
